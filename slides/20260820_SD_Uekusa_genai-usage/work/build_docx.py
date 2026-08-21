#!/usr/bin/env python3
"""植草学園SD研修 ワーク手順書 ── Markdown を1本の .docx に組み上げる。"""
import re, sys, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

S = os.path.dirname(os.path.abspath(__file__))
ACCENT = RGBColor(0x1E, 0x4A, 0x27)
ACCENT2 = RGBColor(0x2F, 0x6B, 0x3A)
JP = 'Yu Gothic'


def set_jp_font(run, name=JP):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    for a in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rf.set(qn(a), name)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+`)')


def add_inline(p, text, base_size=10.5):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); r.bold = True; set_jp_font(r)
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); set_jp_font(r, 'Consolas'); r.font.color.rgb = ACCENT
        else:
            r = p.add_run(part); set_jp_font(r)
        r.font.size = Pt(base_size)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text); r.bold = True; set_jp_font(r)
    r.font.size = Pt({1: 18, 2: 14, 3: 12, 4: 11}.get(level, 10.5))
    r.font.color.rgb = ACCENT if level <= 2 else ACCENT2
    if level == 1:
        pbdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '12'); b.set(qn('w:color'), '1E4A27')
        pbdr.append(b); p._p.get_or_add_pPr().append(pbdr)


def add_table(doc, rows):
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncols); t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j in range(ncols):
            cell = cells[j]; cell.paragraphs[0].text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            add_inline(p, row[j] if j < len(row) else '', base_size=9.5)
            if i == 0:
                shade(cell, 'EAF3EA')
                for r in p.runs:
                    r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def render_markdown(doc, md):
    lines = md.split('\n'); i = 0; in_code = False; buf = []
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('```'):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
                pPr = p._p.get_or_add_pPr()
                pbdr = OxmlElement('w:pBdr')
                for side in ('top', 'left', 'bottom', 'right'):
                    e = OxmlElement('w:' + side)
                    e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '6'); e.set(qn('w:color'), '2F6B3A')
                    pbdr.append(e)
                pPr.append(pbdr)
                shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), 'F5F9F4')
                pPr.append(shd)
                for k, cl in enumerate(buf):
                    r = p.add_run(cl); set_jp_font(r, 'Consolas'); r.font.size = Pt(9.5)
                    if k < len(buf) - 1:
                        r.add_break()
                buf = []; in_code = False
            else:
                in_code = True
            i += 1; continue
        if in_code:
            buf.append(line); i += 1; continue
        if line.strip().startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|[\s:|-]+\|\s*$', lines[i + 1]):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                if re.match(r'^\s*\|[\s:|-]+\|\s*$', lines[i]):
                    i += 1; continue
                rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')]); i += 1
            add_table(doc, rows); continue
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            add_heading(doc, m.group(2).strip(), len(m.group(1))); i += 1; continue
        if re.match(r'^\s*(-{3,}|\*{3,}|_{3,})\s*$', line):
            p = doc.add_paragraph()
            pbdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
            b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6'); b.set(qn('w:color'), 'CCCCCC')
            pbdr.append(b); p._p.get_or_add_pPr().append(pbdr)
            p.paragraph_format.space_after = Pt(8); i += 1; continue
        if line.strip().startswith('>'):
            q = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                q.append(lines[i].strip()[1:].strip()); i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
            pbdr = OxmlElement('w:pBdr'); l = OxmlElement('w:left')
            l.set(qn('w:val'), 'single'); l.set(qn('w:sz'), '18'); l.set(qn('w:color'), '2F6B3A')
            pbdr.append(l); p._p.get_or_add_pPr().append(pbdr)
            for k, bl in enumerate(q):
                if k:
                    p.add_run().add_break()
                add_inline(p, bl, base_size=10)
            continue
        mb = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.*)$', line)
        if mb:
            ordered = bool(re.match(r'^\d+\.$', mb.group(2)))
            p = doc.add_paragraph(style='List Number' if ordered else 'List Bullet')
            p.paragraph_format.left_indent = Cm(0.7 + 0.6 * (len(mb.group(1)) // 2))
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, mb.group(3)); i += 1; continue
        if not line.strip():
            i += 1; continue
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        add_inline(p, line.strip()); i += 1


TITLES = {'A': 'AIと教育のルールを考える', 'B': '既存の授業をAIで進化させる', 'C': 'AI時代の新しい学びを創る',
          'D': '学生とAIの付き合い方を考える', 'E': '教育現場の業務改善を考える', 'F': '広報・情報発信の効率化',
          'G': 'Microsoft 365 Copilot活用', 'H': 'データ活用・集計自動化', 'I': '生成AIによる業務自動化'}


def build_one(md, out):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(1.8)
    sec.left_margin = sec.right_margin = Cm(2.0)
    st = doc.styles['Normal']; st.font.size = Pt(10.5); st.font.name = JP
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), JP)
    fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(); set_jp_font(r); r.font.size = Pt(9)
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE'); fp._p.append(fld)
    render_markdown(doc, md)
    doc.save(out)
    return out


def main():
    outdir = sys.argv[1] if len(sys.argv) > 1 else os.path.join(S, 'docx')
    os.makedirs(outdir, exist_ok=True)
    made = []
    cf = os.path.join(S, 'common-front.md')
    if os.path.exists(cf):
        made.append(build_one(open(cf).read(), os.path.join(outdir, '00_はじめに（全グループ共通）.docx')))
    for k in 'ABCDEFGHI':
        f = os.path.join(S, 'out', f'{k}.md')
        if os.path.exists(f):
            made.append(build_one(open(f).read(), os.path.join(outdir, f'{k}_{TITLES[k]}.docx')))
        else:
            print(f'  !! out/{k}.md なし', file=sys.stderr)
    for m in made:
        print(f'  + {os.path.basename(m)}  ({os.path.getsize(m):,} bytes)')
    print(f'\n{len(made)}ファイルを書き出しました: {outdir}')


if __name__ == '__main__':
    main()
